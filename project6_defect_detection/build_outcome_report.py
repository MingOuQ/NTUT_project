from __future__ import annotations

import csv
import json
import math
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
PROJECT = ROOT / "project6_defect_detection"
PC = PROJECT / "pc_python"
ASSETS = PROJECT / "report_assets"
OUTPUT_DOCX = ROOT / "基於多光譜影像與深度學習之表面瑕疵智慧檢測系統_成果報告.docx"

FONT_REGULAR = Path(r"C:\Windows\Fonts\msjh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\msjhbd.ttc")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(85, 85, 85)
LIGHT_GRAY_HEX = "F2F4F7"
CALL_OUT_HEX = "F4F6F9"
INK = "0B2545"


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    if path.exists():
        return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = ""
        for char in raw_line:
            test = line + char
            if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
                line = test
            else:
                if line:
                    lines.append(line)
                line = char
        if line:
            lines.append(line)
    return lines


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str,
    line_gap: int = 8,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 36)
    line_heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(line_heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) // 2
    for line, h in zip(lines, line_heights):
        w = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + (x2 - x1 - w) // 2, y), line, font=font, fill=fill)
        y += h + line_gap


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str, width: int = 5) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head_len = 20
    for sign in (1, -1):
        a = angle + sign * math.pi * 0.78
        p = (end[0] + head_len * math.cos(a), end[1] + head_len * math.sin(a))
        draw.line([end, p], fill=color, width=width)


def make_system_architecture(path: Path) -> None:
    img = Image.new("RGB", (1800, 1050), "white")
    d = ImageDraw.Draw(img)
    title_font = pil_font(48, True)
    head_font = pil_font(34, True)
    body_font = pil_font(26)
    small_font = pil_font(23)

    d.rectangle([0, 0, 1800, 1050], fill="#F8FAFC")
    d.text((80, 60), "系統架構與資料流", font=title_font, fill="#0B2545")
    d.text((80, 125), "ESP32S 控制節點、ESP32-CAM 影像節點、PC AI 中樞與手機 PWA 形成端到端檢測閉環", font=small_font, fill="#475569")

    boxes = {
        "ctrl": (95, 250, 555, 610),
        "cam": (95, 730, 555, 950),
        "pc": (760, 250, 1240, 610),
        "phone": (1450, 280, 1710, 840),
    }
    colors = {
        "ctrl": ("#DBEAFE", "#2E74B5"),
        "cam": ("#ECFDF5", "#059669"),
        "pc": ("#FFF7ED", "#EA580C"),
        "phone": ("#F1F5F9", "#334155"),
    }
    labels = {
        "ctrl": ("ESP32S 控制節點", "RGB LED 光源控制\nOLED 顯示狀態\n蜂鳴器警報\n實體按鈕觸發\nWi-Fi AP + WebSocket"),
        "cam": ("ESP32-CAM 影像節點", "OV2640 拍攝\nJPEG 壓縮\nIMG:<size> 二進位傳輸\n支援 QVGA/VGA"),
        "pc": ("PC Python AI 中樞", "Serial 協調雙板\nRGB 多光譜擷取\nYOLOv8n 推論\nOpenCV 備援分析\nExcel 報表輸出"),
        "phone": ("手機 PWA 儀表板", "連線 Wi-Fi 熱點\nWebSocket 即時接收\n狀態燈號與音效警報\n歷史紀錄顯示"),
    }

    for key, box in boxes.items():
        fill, outline = colors[key]
        d.rounded_rectangle(box, radius=26, fill=fill, outline=outline, width=4)
        x1, y1, x2, y2 = box
        d.text((x1 + 34, y1 + 28), labels[key][0], font=head_font, fill="#0B2545")
        y = y1 + 95
        for line in labels[key][1].splitlines():
            d.text((x1 + 42, y), line, font=body_font, fill="#1E293B")
            y += 43

    draw_arrow(d, (555, 410), (760, 410), "#2563EB")
    draw_arrow(d, (760, 470), (555, 470), "#2563EB")
    d.text((585, 372), "USB Serial", font=small_font, fill="#1D4ED8")
    d.text((575, 500), "RESULT / MSG", font=small_font, fill="#1D4ED8")

    draw_arrow(d, (555, 840), (760, 540), "#059669")
    draw_arrow(d, (760, 570), (555, 880), "#059669")
    d.text((572, 690), "CAPTURE + JPEG", font=small_font, fill="#047857")

    draw_arrow(d, (555, 330), (1450, 500), "#64748B", 4)
    d.text((960, 305), "Wi-Fi AP / WebSocket 即時狀態", font=small_font, fill="#475569")
    draw_arrow(d, (1240, 395), (1450, 600), "#EA580C", 4)
    d.text((1260, 450), "瑕疵結果同步", font=small_font, fill="#C2410C")

    d.rounded_rectangle((760, 750, 1240, 950), radius=26, fill="#F8FAFC", outline="#94A3B8", width=3)
    d.text((795, 785), "輸出成果", font=head_font, fill="#0B2545")
    for i, line in enumerate(["偵測畫面標註", "Excel 瑕疵座標報表", "模型訓練曲線與驗證圖"]):
        d.text((820, 845 + i * 35), line, font=small_font, fill="#334155")
    draw_arrow(d, (1000, 610), (1000, 750), "#94A3B8", 4)

    img.save(path, quality=95)


def make_inspection_flow(path: Path) -> None:
    img = Image.new("RGB", (1800, 640), "white")
    d = ImageDraw.Draw(img)
    title_font = pil_font(44, True)
    step_font = pil_font(29, True)
    body_font = pil_font(23)
    d.rectangle([0, 0, 1800, 640], fill="#FFFFFF")
    d.text((70, 50), "全自動瑕疵檢測流程", font=title_font, fill="#0B2545")
    steps = [
        ("1", "啟動檢測", "按鈕或空白鍵觸發"),
        ("2", "RGB 打光", "R/G/B 依序切換"),
        ("3", "影像擷取", "ESP32-CAM 回傳 JPEG"),
        ("4", "影像融合", "取出對應通道並合成"),
        ("5", "AI 判讀", "YOLOv8n 定位與分類"),
        ("6", "回饋存檔", "OLED、蜂鳴器、PWA、Excel"),
    ]
    x = 70
    y = 210
    w = 245
    h = 245
    palette = ["#DBEAFE", "#DCFCE7", "#FEF3C7", "#E0F2FE", "#FFE4E6", "#F1F5F9"]
    outline = ["#2563EB", "#16A34A", "#D97706", "#0284C7", "#E11D48", "#475569"]
    for i, (num, head, body) in enumerate(steps):
        bx = (x + i * 285, y, x + i * 285 + w, y + h)
        d.rounded_rectangle(bx, radius=24, fill=palette[i], outline=outline[i], width=4)
        d.ellipse((bx[0] + 22, bx[1] + 22, bx[0] + 78, bx[1] + 78), fill=outline[i])
        draw_centered_text(d, (bx[0] + 22, bx[1] + 22, bx[0] + 78, bx[1] + 78), num, step_font, "white")
        draw_centered_text(d, (bx[0] + 20, bx[1] + 85, bx[2] - 20, bx[1] + 145), head, step_font, "#0B2545")
        draw_centered_text(d, (bx[0] + 22, bx[1] + 150, bx[2] - 22, bx[3] - 10), body, body_font, "#334155")
        if i < len(steps) - 1:
            draw_arrow(d, (bx[2] + 12, y + h // 2), (bx[2] + 40, y + h // 2), "#64748B", 4)
    img.save(path, quality=95)


def fit_image(src: Path, size: tuple[int, int]) -> Image.Image:
    img = Image.open(src).convert("RGB")
    img.thumbnail(size, Image.Resampling.LANCZOS)
    canvas = Image.new("RGB", size, "#FFFFFF")
    canvas.paste(img, ((size[0] - img.width) // 2, (size[1] - img.height) // 2))
    return canvas


def make_multispectral_strip(path: Path) -> None:
    sources = [
        (PC / "test_output" / "panel_red.jpg", "紅光影像"),
        (PC / "test_output" / "panel_green.jpg", "綠光影像"),
        (PC / "test_output" / "panel_blue.jpg", "藍光影像"),
        (PC / "test_output" / "fusion.jpg", "多通道融合 / 檢測輸出"),
    ]
    tile_w, tile_h = 760, 450
    img = Image.new("RGB", (1680, 1120), "#F8FAFC")
    d = ImageDraw.Draw(img)
    title_font = pil_font(44, True)
    cap_font = pil_font(28, True)
    d.text((70, 45), "RGB 多光譜取樣與融合示意", font=title_font, fill="#0B2545")
    positions = [(70, 140), (850, 140), (70, 620), (850, 620)]
    for (src, label), pos in zip(sources, positions):
        tile = fit_image(src, (tile_w, tile_h))
        x, y = pos
        d.rounded_rectangle((x - 8, y - 8, x + tile_w + 8, y + tile_h + 58), radius=18, fill="#FFFFFF", outline="#CBD5E1", width=2)
        img.paste(tile, pos)
        d.text((x + 20, y + tile_h + 14), label, font=cap_font, fill="#1E293B")
    img.save(path, quality=95)


def make_dashboard_mock(path: Path) -> None:
    img = Image.new("RGB", (1250, 900), "#F8FAFC")
    d = ImageDraw.Draw(img)
    title_font = pil_font(42, True)
    head_font = pil_font(28, True)
    body_font = pil_font(22)
    small_font = pil_font(19)

    d.text((70, 55), "手機端 PWA 即時監控介面示意", font=title_font, fill="#0B2545")
    phone = (110, 150, 520, 840)
    d.rounded_rectangle(phone, radius=45, fill="#0F1016", outline="#334155", width=7)
    screen = (135, 185, 495, 805)
    d.rounded_rectangle(screen, radius=32, fill="#111827")
    d.text((165, 220), "表面瑕疵監控 App", font=head_font, fill="#E2E8F0")
    d.rounded_rectangle((165, 285, 465, 500), radius=26, fill="#3F0F22", outline="#FF0055", width=3)
    d.ellipse((272, 320, 358, 406), fill="#FF0055")
    draw_centered_text(d, (185, 415, 445, 465), "DEFECT", head_font, "#F8FAFC")
    draw_centered_text(d, (185, 462, 445, 492), "偵測到 3 個瑕疵", small_font, "#FECACA")
    d.rounded_rectangle((165, 525, 465, 610), radius=18, fill="#172554", outline="#00D2FF", width=2)
    d.text((190, 548), "WebSocket: 連線中", font=small_font, fill="#BAE6FD")
    d.rounded_rectangle((165, 635, 465, 765), radius=18, fill="#1F2937", outline="#475569", width=2)
    for i, line in enumerate(["18:31:38 square 85%", "18:31:38 circle 84%", "18:31:38 line_h 76%"]):
        d.text((190, 660 + i * 34), line, font=small_font, fill="#E5E7EB")

    cards = [
        ("即時推播", "ESP32S 以 WebSocket 將 SCANNING、DEFECT、DONE 狀態送到手機。"),
        ("現場警報", "OLED 顯示瑕疵數量與類型，蜂鳴器提供聲音警示。"),
        ("無外網運作", "控制板建立 Defect_Detector_WiFi 熱點，手機直接連線即可監控。"),
    ]
    for i, (head, body) in enumerate(cards):
        y = 220 + i * 170
        d.rounded_rectangle((620, y, 1160, y + 125), radius=18, fill="#FFFFFF", outline="#CBD5E1", width=2)
        d.text((650, y + 25), head, font=head_font, fill="#0B2545")
        for j, line in enumerate(wrap_text(d, body, body_font, 455)):
            d.text((650, y + 67 + j * 31), line, font=body_font, fill="#334155")
    img.save(path, quality=95)


def read_training_metrics() -> dict[str, float | int]:
    csv_path = PC / "runs" / "detect" / "mini_led_defect" / "results.csv"
    with csv_path.open(newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))

    def val(row: dict[str, str], key: str) -> float:
        return float(row.get(key) or 0)

    last = rows[-1]
    best_m50 = max(rows, key=lambda r: val(r, "metrics/mAP50(B)"))
    best_m5095 = max(rows, key=lambda r: val(r, "metrics/mAP50-95(B)"))
    return {
        "epochs": len(rows),
        "last_precision": val(last, "metrics/precision(B)"),
        "last_recall": val(last, "metrics/recall(B)"),
        "last_map50": val(last, "metrics/mAP50(B)"),
        "last_map5095": val(last, "metrics/mAP50-95(B)"),
        "best_map50": val(best_m50, "metrics/mAP50(B)"),
        "best_map50_epoch": int(float(best_m50["epoch"])),
        "best_map5095": val(best_m5095, "metrics/mAP50-95(B)"),
        "best_map5095_epoch": int(float(best_m5095["epoch"])),
    }


def make_metrics_card(path: Path, metrics: dict[str, float | int]) -> None:
    img = Image.new("RGB", (1600, 760), "#F8FAFC")
    d = ImageDraw.Draw(img)
    title_font = pil_font(42, True)
    metric_font = pil_font(48, True)
    head_font = pil_font(24, True)
    body_font = pil_font(22)
    d.text((70, 45), "模型訓練與驗證摘要", font=title_font, fill="#0B2545")

    cards = [
        ("訓練紀錄", f"{metrics['epochs']} epochs", "results.csv 最後紀錄"),
        ("最佳 mAP50", f"{metrics['best_map50']:.3f}", f"出現在 epoch {metrics['best_map50_epoch']}"),
        ("最佳 mAP50-95", f"{metrics['best_map5095']:.3f}", f"出現在 epoch {metrics['best_map5095_epoch']}"),
        ("最後 Precision / Recall", f"{metrics['last_precision']:.3f} / {metrics['last_recall']:.3f}", "模型已能穩定定位主要瑕疵"),
    ]
    for i, (head, value, note) in enumerate(cards):
        x = 70 + i * 375
        d.rounded_rectangle((x, 150, x + 330, 350), radius=24, fill="#FFFFFF", outline="#CBD5E1", width=2)
        d.text((x + 28, 178), head, font=head_font, fill="#334155")
        d.text((x + 28, 225), value, font=metric_font, fill="#2E74B5")
        d.text((x + 28, 305), note, font=body_font, fill="#64748B")

    csv_path = PC / "runs" / "detect" / "mini_led_defect" / "results.csv"
    rows = list(csv.DictReader(csv_path.open(newline="", encoding="utf-8")))
    epochs = [int(float(r["epoch"])) for r in rows]
    map50 = [float(r["metrics/mAP50(B)"]) for r in rows]
    box_loss = [float(r["val/box_loss"]) for r in rows]

    chart = (120, 440, 1480, 680)
    d.rounded_rectangle((70, 400, 1530, 720), radius=22, fill="#FFFFFF", outline="#CBD5E1", width=2)
    d.text((110, 420), "趨勢：mAP50 上升，Validation box loss 下降", font=head_font, fill="#0B2545")
    left, top, right, bottom = chart
    d.line((left, bottom, right, bottom), fill="#94A3B8", width=2)
    d.line((left, top, left, bottom), fill="#94A3B8", width=2)

    def plot_line(values: list[float], color: str, min_v: float | None = None, max_v: float | None = None) -> None:
        min_v = min(values) if min_v is None else min_v
        max_v = max(values) if max_v is None else max_v
        pts = []
        for e, v in zip(epochs, values):
            x = left + (e - min(epochs)) / max(1, (max(epochs) - min(epochs))) * (right - left)
            y = bottom - (v - min_v) / max(0.0001, max_v - min_v) * (bottom - top)
            pts.append((x, y))
        d.line(pts, fill=color, width=5)

    plot_line(map50, "#2563EB", 0, 0.85)
    loss_scaled = [1 / (1 + v) for v in box_loss]
    plot_line(loss_scaled, "#EA580C", min(loss_scaled), max(loss_scaled))
    d.rounded_rectangle((1110, 420, 1465, 492), radius=16, fill="#F8FAFC", outline="#CBD5E1")
    d.line((1135, 445, 1190, 445), fill="#2563EB", width=5)
    d.text((1205, 431), "mAP50", font=body_font, fill="#334155")
    d.line((1320, 445, 1375, 445), fill="#EA580C", width=5)
    d.text((1390, 431), "val loss 趨勢", font=body_font, fill="#334155")
    img.save(path, quality=95)


def latest_excel_rows() -> tuple[str, list[tuple]]:
    files = sorted((PC / "inspection_reports").glob("*.xlsx"))
    if not files:
        return "", []
    latest = files[-1]
    wb = load_workbook(latest, data_only=True)
    ws = wb.active
    rows = []
    for row in ws.iter_rows(min_row=4, max_row=min(ws.max_row, 9), values_only=True):
        rows.append(tuple("" if v is None else v for v in row))
    return latest.name, rows


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=bold, color=color)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "D0D7DE") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths: Iterable[float]) -> None:
    widths_in = list(widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_in):
                cell.width = Inches(widths_in[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def style_table(table, widths: Iterable[float], header: bool = True) -> None:
    set_table_geometry(table, widths)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                set_paragraph_font(p, size=10.3)
            if header and i == 0:
                shade_cell(cell, LIGHT_GRAY_HEX)
                for p in cell.paragraphs:
                    set_paragraph_font(p, bold=True, color=NAVY)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=GRAY)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    set_paragraph_font(p)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    set_paragraph_font(p)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def set_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    header = section.header.paragraphs[0]
    header.text = "表面瑕疵智慧檢測系統成果報告"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_font(header, size=9.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "基於多光譜影像與深度學習之表面瑕疵智慧檢測系統"
    set_paragraph_font(footer, size=9, color=GRAY)


def add_cover(doc: Document, assets: dict[str, Path], metrics: dict[str, float | int]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("成果報告")
    set_run_font(r, size=15, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("基於多光譜影像與深度學習之\n表面瑕疵智慧檢測系統")
    set_run_font(r, size=22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESP32S / ESP32-CAM / Python YOLOv8 自動化檢測整合")
    set_run_font(r, size=13, color=GRAY)

    add_figure(doc, assets["architecture"], "圖 1  系統架構總覽：控制、拍攝、AI 判讀與行動監控的資料流", width=5.35)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("課程名稱", "微處理機及自動控制應用實務"),
        ("專題名稱", "基於機器視覺與光電感測之表面瑕疵智慧檢測系統"),
        ("使用平台", "ESP32S、ESP32-CAM、Python、YOLOv8、OpenCV、WebSocket、PWA"),
        ("整理日期", "2026 年 5 月 23 日"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label, bold=True, color=NAVY)
        set_cell_text(row.cells[1], value)
        shade_cell(row.cells[0], LIGHT_GRAY_HEX)
    style_table(table, [1.4, 5.1], header=False)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"模型最佳 mAP50 約 {metrics['best_map50']:.3f}，可辨識 point、square、circle、line_h、line_v 五類瑕疵。")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    doc.add_page_break()


def add_summary(doc: Document) -> None:
    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "本專題設計並實作一套結合多光譜影像分析、YOLOv8 深度學習物件偵測與邊緣運算概念的自動化表面瑕疵檢測系統。"
        "系統以 ESP32S 作為控制節點，負責 RGB LED 光源切換、OLED 顯示、蜂鳴器警報與 Wi-Fi AP / WebSocket 通訊；"
        "ESP32-CAM 則作為影像擷取節點，依 PC 端指令拍攝並回傳 JPEG 影像。PC 端 Python 程式協調兩個微控制器，"
        "依序取得紅、綠、藍三色照明下的影像，透過多通道融合突顯待測物表面缺陷，再使用 YOLOv8-nano 模型進行瑕疵定位與分類。"
    )
    add_body(
        doc,
        "本系統可辨識 point、square、circle、line_h、line_v 五種 Mini LED 面板常見瑕疵型態，並支援 OpenCV 傳統影像處理作為備援。"
        "檢測結果可即時回傳至 ESP32S，於 OLED 與蜂鳴器端產生現場回饋，也能透過手機 PWA Dashboard 接收狀態推播與警示。"
        "此外，每次檢測皆可匯出 Excel 報表，記錄瑕疵座標、尺寸、類型與信心度，方便後續品質追蹤。"
        "整體成果顯示，本專題已完成從硬體控制、影像擷取、AI 判讀、行動監控到報表輸出的完整閉環。"
    )


def add_intro(doc: Document) -> None:
    add_heading(doc, "1. 前言", 1)
    add_body(
        doc,
        "在電子製造、顯示面板與精密零件生產中，表面瑕疵會直接影響產品良率、可靠度與客戶接受度。"
        "傳統人工目視檢測雖然彈性高，但容易受到疲勞、環境光、檢測人員經驗差異影響，導致漏檢或誤判。"
        "工業級自動光學檢測設備可以改善這些問題，但設備成本高、系統封閉且客製化不易，對課程專題或中小規模應用而言進入門檻偏高。"
    )
    add_body(
        doc,
        "因此，本專題以低成本嵌入式硬體與開源 AI 工具為基礎，實作一套可展示完整 AOI 概念的智慧檢測系統。"
        "此系統並非只做單張影像判讀，而是將光源控制、相機擷取、AI 推論、現場警示、手機監控與資料匯出串成一個可操作流程，"
        "用課程可取得的硬體材料驗證智慧製造應用的可行性。"
    )


def add_background(doc: Document, assets: dict[str, Path]) -> None:
    add_heading(doc, "2. 研究背景與動機", 1)
    add_heading(doc, "2.1 產業痛點", 2)
    add_body(
        doc,
        "Mini LED 或 Micro LED 面板上具有大量微小發光單元，製程中可能出現暗點、異物、刮痕、線狀裂痕或局部凹陷。"
        "這些缺陷尺寸小、對比低，若僅以單一白光照明檢測，某些瑕疵可能因反射角度或材料顏色接近背景而不明顯。"
        "本專題因此導入 RGB 三色光源，利用不同波長與缺陷表面反射差異，提高影像特徵可分辨性。"
    )

    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["瑕疵類型", "外觀特徵", "可能影響"]):
        set_cell_text(cell, text, bold=True, color=NAVY)
    defect_rows = [
        ("point", "微小點狀髒污、暗點或亮點", "影響顯示均勻度"),
        ("square", "方塊狀異物或局部突起", "可能造成短路、遮光或區域失效"),
        ("circle", "圓形凹陷、氣泡或壓痕", "影響結構完整性與發光均勻性"),
        ("line_h", "水平方向刮痕或裂痕", "影響導線或表面連續性"),
        ("line_v", "垂直方向刮痕或裂痕", "同樣可能造成走線或視覺缺陷"),
    ]
    for values in defect_rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)
    style_table(table, [1.3, 3.0, 2.2])

    add_figure(doc, assets["dataset_contact"], "圖 2  自建 Mini LED 瑕疵資料集樣本，包含多種形態與位置的缺陷", width=4.65)

    add_heading(doc, "2.2 研究動機", 2)
    add_body(
        doc,
        "本專題的動機有三點：第一，以 ESP32 系列模組降低 AOI 系統建置門檻；第二，驗證 RGB 多光譜打光在瑕疵顯影上的效果；"
        "第三，把 AI 推論與 IoT 即時監控整合到同一個操作流程中，使系統不只會辨識，也能把結果即時呈現給操作者並保存檢測紀錄。"
    )


def add_goals(doc: Document) -> None:
    add_heading(doc, "3. 研究目標", 1)
    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ["目標", "實作內容", "對應成果"]):
        set_cell_text(cell, text, bold=True, color=NAVY)
    rows = [
        ("三節點硬體架構", "ESP32S 控制、ESP32-CAM 拍攝、PC 端推論分工", "降低單一節點負擔並提升可維護性"),
        ("多光譜檢測流程", "R/G/B 三色 LED 依序照射並擷取影像", "提升不同瑕疵在影像中的對比"),
        ("AI 瑕疵辨識", "以 YOLOv8n 訓練五類缺陷偵測模型", "輸出瑕疵座標、類別與信心度"),
        ("行動端監控", "ESP32S 建立 Wi-Fi AP，手機以 PWA 接收 WebSocket 訊息", "現場可即時掌握檢測狀態與警示"),
        ("報表自動化", "使用 openpyxl 匯出 Excel 檢測報告", "保留品質追蹤與後續分析資料"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)
    style_table(table, [1.45, 3.0, 2.05])
    add_body(
        doc,
        "上述目標共同構成本專題的核心：不是單點功能展示，而是從感測到回饋、從模型到資料紀錄的一套完整檢測流程。"
    )


def add_principles(doc: Document, assets: dict[str, Path]) -> None:
    add_heading(doc, "4. 原理說明", 1)
    add_heading(doc, "4.1 系統架構與資料流", 2)
    add_body(
        doc,
        "系統採用三節點分散式架構。ESP32S 負責控制 RGB 光源、OLED、蜂鳴器、按鈕與手機端 WebSocket；"
        "ESP32-CAM 負責拍攝影像並將 JPEG 資料以 Serial 傳回 PC；PC 端 Python 程式則是流程協調與 AI 推論核心。"
        "這樣的設計讓即時控制與高運算量影像分析各自落在適合的硬體上。"
    )
    add_figure(doc, assets["flow"], "圖 3  全自動檢測流程：啟動、打光、拍攝、融合、AI 判讀與結果回饋", width=6.3)

    add_heading(doc, "4.2 多光譜影像融合", 2)
    add_body(
        doc,
        "待測物在紅、綠、藍不同波長光線照射下，瑕疵區域的反射特性會與正常區域產生差異。"
        "PC 端會從紅光影像取 R 通道、綠光影像取 G 通道、藍光影像取 B 通道，再進行融合。"
        "專案程式中使用 element-wise minimum 的方式產生融合灰階圖，其概念可表示為："
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("F(x, y) = min(Rred(x, y), Ggreen(x, y), Bblue(x, y))")
    set_run_font(r, size=12, bold=True, color=NAVY)
    add_body(
        doc,
        "取最小值可使在任一光譜下呈現暗化或反射異常的缺陷更容易被保留，搭配後續 YOLO 或 OpenCV 演算法進行定位。"
    )
    add_figure(doc, assets["multispectral"], "圖 4  紅、綠、藍光影像與融合結果示意", width=5.95)

    add_heading(doc, "4.3 YOLOv8 與 OpenCV 雙模式偵測", 2)
    add_body(
        doc,
        "主要辨識方法為 YOLOv8-nano 物件偵測模型。資料集採 YOLO detection 標註格式，模型輸出每個瑕疵的邊界框、類別與信心度。"
        "若 AI 權重不存在或 ultralytics 套件不可用，PC 程式會退回 OpenCV 模式，以 Gaussian blur、adaptive threshold、morphology 與輪廓分析找出疑似瑕疵。"
        "雙模式設計讓系統在模型尚未訓練完成時仍可進行基礎測試。"
    )

    doc.add_page_break()
    add_heading(doc, "4.4 通訊與即時回饋", 2)
    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ["通訊方向", "指令或資料", "功能"]):
        set_cell_text(cell, text, bold=True, color=NAVY)
    rows = [
        ("PC -> ESP32S", "LED:R / LED:G / LED:B", "切換三色 LED 光源"),
        ("PC -> ESP32-CAM", "CAPTURE", "觸發相機拍照"),
        ("ESP32-CAM -> PC", "IMG:<size> + JPEG bytes", "回傳影像資料"),
        ("PC -> ESP32S", "RESULT:DEFECT,...", "傳回瑕疵座標、類型與信心度"),
        ("ESP32S -> 手機", "WebSocket JSON / 狀態訊息", "即時更新 PWA Dashboard"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)
    style_table(table, [1.45, 2.25, 2.8])


def add_results(doc: Document, assets: dict[str, Path], metrics: dict[str, float | int]) -> None:
    add_heading(doc, "5. 結果討論", 1)
    add_heading(doc, "5.1 系統功能完成度", 2)
    add_body(
        doc,
        "本專題已完成控制節點、相機節點與 PC 推論節點的整合，並能透過按鈕或 OpenCV 視窗快捷鍵啟動檢測。"
        "系統支援單面板與 3x3 拼接面板模式，檢測完成後會把結果同步回 ESP32S，並在手機端與 Excel 報表中留下紀錄。"
    )
    add_figure(doc, assets["nine_panel"], "圖 5  3x3 拼接面板模式測試結果，系統可對各區塊標示 OK / NG", width=5.8)

    add_heading(doc, "5.2 AI 訓練成果", 2)
    add_body(
        doc,
        f"目前訓練紀錄共 {metrics['epochs']} 筆 epoch 紀錄。依 results.csv 統計，最佳 mAP50 約為 {metrics['best_map50']:.3f}，"
        f"出現在 epoch {metrics['best_map50_epoch']}；最佳 mAP50-95 約為 {metrics['best_map5095']:.3f}，出現在 epoch {metrics['best_map5095_epoch']}。"
        f"最後一筆紀錄的 precision 約 {metrics['last_precision']:.3f}、recall 約 {metrics['last_recall']:.3f}，顯示模型已可穩定辨識資料集中主要瑕疵形態。"
    )
    add_figure(doc, assets["metrics"], "圖 6  由訓練紀錄整理之模型指標摘要", width=6.25)
    add_figure(doc, assets["training_plot"], "圖 7  YOLOv8 訓練曲線：loss 下降且 mAP 指標逐步提升", width=6.25)
    add_figure(doc, assets["confusion"], "圖 8  模型混淆矩陣，可觀察各類瑕疵分類狀況", width=5.55)

    add_heading(doc, "5.3 檢測資料與報表輸出", 2)
    report_name, excel_rows = latest_excel_rows()
    add_body(
        doc,
        f"PC 程式會在每次檢測後自動於 inspection_reports 資料夾產生 Excel 報表。"
        f"目前資料夾內已有多份測試報告，最新範例為 {report_name}，內容包含中心座標、寬高、瑕疵類型、信心度與面積。"
    )
    if excel_rows:
        table = doc.add_table(rows=1, cols=8)
        for cell, text in zip(table.rows[0].cells, excel_rows[0]):
            set_cell_text(cell, str(text), bold=True, color=NAVY)
        for values in excel_rows[1:]:
            cells = table.add_row().cells
            for cell, value in zip(cells, values):
                if isinstance(value, float):
                    value = f"{value:.2f}"
                set_cell_text(cell, str(value))
        style_table(table, [0.45, 0.75, 0.75, 0.65, 0.65, 1.25, 0.85, 0.75])

    add_heading(doc, "5.4 行動監控與現場可用性", 2)
    add_body(
        doc,
        "ESP32S 可自行建立 Wi-Fi AP 熱點，手機連線後開啟 PWA Dashboard 即可接收檢測狀態。"
        "此設計不依賴外部路由器或校園網路，展示環境部署較容易；同時 OLED 與蜂鳴器可提供機台端即時回饋，"
        "手機端則提供較完整的狀態視覺化與歷史訊息。"
    )
    add_figure(doc, assets["dashboard"], "圖 9  手機端 PWA 即時監控介面示意", width=5.8)

    doc.add_page_break()
    add_heading(doc, "5.5 問題與限制", 2)
    table = doc.add_table(rows=1, cols=2)
    for cell, text in zip(table.rows[0].cells, ["觀察到的問題", "改善方式"]):
        set_cell_text(cell, text, bold=True, color=NAVY)
    rows = [
        ("ESP32-CAM 供電不足時可能 brownout 或拍照失敗", "改用 5V 穩定供電，並降低解析度至 QVGA 以提高穩定性"),
        ("模擬資料與真實拍攝影像存在 domain gap", "持續使用系統收集真實影像，透過 labelImg 標註後再 fine-tuning"),
        ("Serial Port 可能被 Arduino IDE 佔用", "執行 Python 前關閉 Serial Monitor，並在程式中控制 DTR/RTS"),
        ("資料集規模仍偏小", "擴充不同材質、角度、亮度與瑕疵尺寸的樣本，提高泛化能力"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)
    style_table(table, [3.0, 3.5])


def add_conclusion(doc: Document) -> None:
    add_heading(doc, "6. 結論", 1)
    add_body(
        doc,
        "本專題完成一套以 ESP32S、ESP32-CAM 與 PC Python AI 模組構成的表面瑕疵智慧檢測系統。"
        "系統能自動控制 RGB 三色光源、擷取多光譜影像、進行 YOLOv8 瑕疵偵測，並將檢測結果同步至 OLED、蜂鳴器、手機 PWA 與 Excel 報表。"
        "從課程成果角度來看，本系統同時涵蓋微控制器控制、感測資料傳輸、影像處理、深度學習模型訓練、無線通訊與資料視覺化，具有完整的整合展示價值。"
    )
    add_body(
        doc,
        "結果顯示，低成本硬體搭配合理的影像流程與 AI 模型，能建立具雛形實用性的 AOI 檢測平台。"
        "未來若要提升到更接近工業應用的程度，可持續擴充真實資料集、改善光學機構與相機固定方式、導入更高解析度鏡頭，"
        "並評估將模型轉換為可在邊緣端執行的格式。整體而言，本專題已達成智慧瑕疵檢測系統的原型開發目標，"
        "也展示了嵌入式控制與 AI 視覺在自動化品質檢測中的應用潛力。"
    )

    table = doc.add_table(rows=1, cols=2)
    for cell, text in zip(table.rows[0].cells, ["已完成成果", "後續優化方向"]):
        set_cell_text(cell, text, bold=True, color=NAVY)
    rows = [
        ("三節點硬體整合與 Serial 通訊", "提高相機解析度與固定光學機構"),
        ("RGB 多光譜拍攝與影像融合", "擴充真實瑕疵資料集"),
        ("YOLOv8n 五類瑕疵偵測模型", "加入更多量化指標與測試樣本"),
        ("PWA 即時監控與 Excel 報表", "評估邊緣 AI 或 TFLite 部署"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)
    style_table(table, [3.15, 3.35])


def make_report() -> Path:
    ASSETS.mkdir(exist_ok=True)
    metrics = read_training_metrics()

    assets = {
        "architecture": ASSETS / "system_architecture.png",
        "flow": ASSETS / "inspection_flow.png",
        "multispectral": ASSETS / "multispectral_strip.png",
        "dashboard": ASSETS / "dashboard_mock.png",
        "metrics": ASSETS / "metrics_card.png",
        "dataset_contact": ROOT / "mini_led_defect_dataset_rgb_label_v2" / "sample_contact_sheet.png",
        "nine_panel": PC / "test_output" / "9panel_result.jpg",
        "training_plot": PC / "runs" / "detect" / "mini_led_defect" / "results.png",
        "confusion": PC / "runs" / "detect" / "mini_led_defect" / "confusion_matrix.png",
    }
    make_system_architecture(assets["architecture"])
    make_inspection_flow(assets["flow"])
    make_multispectral_strip(assets["multispectral"])
    make_dashboard_mock(assets["dashboard"])
    make_metrics_card(assets["metrics"], metrics)

    doc = Document()
    set_styles(doc)
    add_cover(doc, assets, metrics)
    add_summary(doc)
    add_intro(doc)
    add_background(doc, assets)
    doc.add_page_break()
    add_goals(doc)
    add_principles(doc, assets)
    add_results(doc, assets, metrics)
    doc.add_page_break()
    add_conclusion(doc)

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(make_report())
